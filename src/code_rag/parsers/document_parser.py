"""
Document parser for Word (.docx) files.

Extracts text, images, and structure from corporate documentation (SOP, policies, etc).
"""

import io
import hashlib
from pathlib import Path
from typing import List, Dict, Any, Optional
from dataclasses import dataclass, field

from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import Table
from PIL import Image

from src.logger import get_logger


logger = get_logger(__name__)


@dataclass
class DocumentSection:
    """Section of a document."""
    title: str = ""
    level: int = 0  # Heading level (0 = no heading, 1 = H1, 2 = H2, etc)
    content: str = ""
    images: List[Dict[str, Any]] = field(default_factory=list)
    tables: List[Dict[str, Any]] = field(default_factory=list)
    position: int = 0  # Position in document


@dataclass
class ParsedDocument:
    """Parsed document result."""
    file_path: Path
    title: str
    full_text: str
    sections: List[DocumentSection] = field(default_factory=list)
    images: List[Dict[str, Any]] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)
    errors: List[str] = field(default_factory=list)


class DocumentParser:
    """
    Parser for .docx documents.

    Extracts:
    - Text content with structure
    - Images with positions
    - Tables
    - Metadata (title, author, dates)
    """

    def __init__(self, images_dir: Optional[Path] = None):
        """
        Initialize document parser.

        Args:
            images_dir: Directory to save extracted images (default: data/document_images)
        """
        if images_dir is None:
            images_dir = Path("data/document_images")

        self.images_dir = images_dir
        self.images_dir.mkdir(parents=True, exist_ok=True)

    def parse_file(self, file_path: Path) -> ParsedDocument:
        """
        Parse a .docx file.

        Args:
            file_path: Path to .docx file

        Returns:
            ParsedDocument with extracted content
        """
        try:
            logger.info(f"Parsing document: {file_path}")

            # Load document
            doc = Document(file_path)

            # Extract metadata
            metadata = self._extract_metadata(doc)

            # Extract title
            title = self._extract_title(doc, file_path)

            # Extract content sections
            sections = self._extract_sections(doc, file_path)

            # Extract all images
            all_images = self._extract_images(doc, file_path)

            # Build full text
            full_text = "\n\n".join([
                f"{'#' * (s.level or 1)} {s.title}\n{s.content}" if s.title else s.content
                for s in sections
            ])

            logger.info(
                f"Parsed {file_path.name}: {len(sections)} sections, "
                f"{len(all_images)} images, {len(full_text)} chars"
            )

            return ParsedDocument(
                file_path=file_path,
                title=title,
                full_text=full_text,
                sections=sections,
                images=all_images,
                metadata=metadata,
                errors=[]
            )

        except Exception as e:
            logger.error(f"Failed to parse {file_path}: {e}")
            return ParsedDocument(
                file_path=file_path,
                title=file_path.stem,
                full_text="",
                sections=[],
                images=[],
                metadata={},
                errors=[str(e)]
            )

    def _extract_metadata(self, doc: Document) -> Dict[str, Any]:
        """Extract document metadata (properties)."""
        metadata = {}

        try:
            core_props = doc.core_properties

            metadata['author'] = core_props.author or ""
            metadata['title'] = core_props.title or ""
            metadata['subject'] = core_props.subject or ""
            metadata['created'] = core_props.created.isoformat() if core_props.created else ""
            metadata['modified'] = core_props.modified.isoformat() if core_props.modified else ""
            metadata['last_modified_by'] = core_props.last_modified_by or ""

        except Exception as e:
            logger.warning(f"Failed to extract metadata: {e}")

        return metadata

    def _extract_title(self, doc: Document, file_path: Path) -> str:
        """Extract document title."""
        # Try from metadata
        if doc.core_properties.title:
            return doc.core_properties.title

        # Try first heading
        for para in doc.paragraphs:
            if para.style.name.startswith('Heading') or para.text.strip():
                title = para.text.strip()
                if title:
                    return title

        # Fallback to filename
        return file_path.stem

    def _extract_sections(self, doc: Document, file_path: Path) -> List[DocumentSection]:
        """
        Extract document sections based on structure.

        Sections are created based on headings or paragraph breaks.
        """
        sections = []
        current_section = DocumentSection(position=0)
        position = 0

        for block in self._iter_block_items(doc):
            position += 1

            if isinstance(block, Paragraph):
                para = block

                # Check if heading
                style_name = para.style.name
                if style_name.startswith('Heading'):
                    # Save previous section if has content
                    if current_section.content or current_section.images:
                        sections.append(current_section)

                    # Start new section
                    level = self._extract_heading_level(style_name)
                    current_section = DocumentSection(
                        title=para.text.strip(),
                        level=level,
                        content="",
                        position=position
                    )
                else:
                    # Regular paragraph - add to current section
                    text = para.text.strip()
                    if text:
                        if current_section.content:
                            current_section.content += "\n\n"
                        current_section.content += text

                # Check for inline images
                inline_images = self._extract_paragraph_images(para, file_path, position)
                current_section.images.extend(inline_images)

            elif isinstance(block, Table):
                # Extract table
                table_data = self._extract_table(block)
                current_section.tables.append({
                    'position': position,
                    'data': table_data
                })

        # Add last section
        if current_section.content or current_section.images or current_section.tables:
            sections.append(current_section)

        # If no sections created (flat document), create one section
        if not sections:
            full_text = "\n\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            sections.append(DocumentSection(
                title=self._extract_title(doc, file_path),
                content=full_text,
                position=0
            ))

        return sections

    def _iter_block_items(self, doc: Document):
        """
        Iterate through document blocks (paragraphs and tables) in order.

        Yields:
            Paragraph or Table objects in document order
        """
        parent = doc.element.body
        for child in parent.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, doc)
            elif isinstance(child, CT_Tbl):
                yield Table(child, doc)

    def _extract_heading_level(self, style_name: str) -> int:
        """Extract heading level from style name."""
        if 'Heading 1' in style_name:
            return 1
        elif 'Heading 2' in style_name:
            return 2
        elif 'Heading 3' in style_name:
            return 3
        elif 'Heading 4' in style_name:
            return 4
        elif 'Heading' in style_name:
            # Try to extract number
            try:
                level = int(''.join(filter(str.isdigit, style_name)))
                return level
            except:
                return 1
        return 0

    def _extract_paragraph_images(
        self,
        para: Paragraph,
        file_path: Path,
        position: int
    ) -> List[Dict[str, Any]]:
        """Extract images from a paragraph."""
        images = []

        try:
            for run in para.runs:
                # Check for inline shapes (images)
                if 'graphic' in run._element.xml:
                    for rel in run.part.rels.values():
                        if "image" in rel.target_ref:
                            image_data = rel.target_part.blob

                            # Save image
                            image_info = self._save_image(
                                image_data,
                                file_path,
                                position
                            )

                            if image_info:
                                images.append(image_info)

        except Exception as e:
            logger.warning(f"Failed to extract images from paragraph: {e}")

        return images

    def _extract_images(self, doc: Document, file_path: Path) -> List[Dict[str, Any]]:
        """Extract all images from document."""
        images = []

        try:
            # Get all image relationships
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_data = rel.target_part.blob

                    # Save image
                    image_info = self._save_image(
                        image_data,
                        file_path,
                        position=len(images)
                    )

                    if image_info:
                        images.append(image_info)

        except Exception as e:
            logger.warning(f"Failed to extract images: {e}")

        return images

    def _save_image(
        self,
        image_data: bytes,
        file_path: Path,
        position: int
    ) -> Optional[Dict[str, Any]]:
        """Save image to disk and return metadata."""
        try:
            # Generate unique filename based on content hash
            image_hash = hashlib.md5(image_data).hexdigest()[:12]
            doc_name = file_path.stem

            # Determine format
            image = Image.open(io.BytesIO(image_data))
            format_lower = image.format.lower() if image.format else 'png'

            # Save path
            image_filename = f"{doc_name}_{position}_{image_hash}.{format_lower}"
            image_path = self.images_dir / image_filename

            # Save image
            image.save(image_path)

            return {
                'path': str(image_path),
                'filename': image_filename,
                'size': len(image_data),
                'width': image.width,
                'height': image.height,
                'format': image.format,
                'position': position
            }

        except Exception as e:
            logger.warning(f"Failed to save image: {e}")
            return None

    def _extract_table(self, table: Table) -> List[List[str]]:
        """Extract table data as list of rows."""
        table_data = []

        try:
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data.append(row_data)

        except Exception as e:
            logger.warning(f"Failed to extract table: {e}")

        return table_data
